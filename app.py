import streamlit as st
import PyPDF2
import docx
import io
from groq import Groq
import json
import re
import os
from dotenv import load_dotenv
from deep_translator import GoogleTranslator
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches

# Load environment variables
load_dotenv()

# Initialize Groq client
llm = Groq(api_key=os.getenv("GROQ_API_KEY"))

# Function to extract text from PDF
def text_from_pdf(file):
    reader = PyPDF2.PdfReader(file)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + " "
    return text

# Function to extract text from DOCX
def text_from_docx(file):
    doc = docx.Document(file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + " "
    return text

# Function to extract text from text file
def text_from_text(file):
    try:
        return file.read().decode("utf-8")
    except UnicodeDecodeError:
        # Try different encodings if utf-8 fails
        try:
            return file.read().decode("latin-1")
        except:
            st.error("Could not decode file. Please check encoding.")
            return ""

# Function to extract structured data using Groq
def data_extraction(text):
    # If text is too long, chunk it
    max_chunk_size = 8000  # Adjust based on model's context window
    
    if len(text) > max_chunk_size:
        chunks = [text[i:i+max_chunk_size] for i in range(0, len(text), max_chunk_size)]
        all_results = []
        for i, chunk in enumerate(chunks):
            with st.status(f"Processing chunk {i+1}/{len(chunks)}..."):
                chunk_result = process_chunk(chunk)
                if chunk_result:
                    all_results.append(chunk_result)
        
        # Merge results
        merged_result = merge_results(all_results)
        return json.dumps(merged_result, indent=2)
    else:
        # Add json.dumps here to ensure we always return a string
        result = process_chunk(text)
        if result:
            return json.dumps(result, indent=2)
        return None

def process_chunk(chunk):
    prompt = """
    You are an advanced medical information extraction specialist. Your task is to accurately extract and categorize all relevant medical entities from the provided text.

    Text: {chunk}
    
    Extract and categorize entities into the following categories:
    1. "DISEASE": Diseases, conditions, disorders, and clinical diagnoses.
    2. "MEDICINE": Medications, drugs, active ingredients, and pharmaceutical products.
    3. "TREATMENT": Medical procedures, therapies, interventions, and surgeries.
    4. "TEST": Laboratory tests, imaging studies, and diagnostic procedures.
    5. "ANATOMY": Body parts, organs, anatomical structures.
    6. "SYMPTOMS": Signs, symptoms, clinical manifestations, and observable conditions.

    Return the output strictly as a well-formed JSON object with the specified categories as keys and arrays of unique entities as values. 
    Avoid any additional text or explanation. Maintain case sensitivity and eliminate duplicates.

    Example output format:
    {{
        "DISEASE": ["hypertension", "type 2 diabetes"],
        "MEDICINE": ["metformin", "lisinopril"],
        "TREATMENT": ["physical therapy", "cognitive behavioral therapy"],
        "TEST": ["blood glucose test", "MRI"],
        "ANATOMY": ["pancreas", "kidney"],
        "SYMPTOMS": ["headache", "fatigue"]
    }}
"""

    try:
        response = llm.chat.completions.create(
            messages=[{"role": "user", "content": prompt.format(chunk=chunk)}],
            model="mixtral-8x7b-32768"
        )
        return json_from_text(response.choices[0].message.content)
    except Exception as e:
        st.error(f"API Error: {str(e)}")
        return None

def merge_results(results_list):
    merged = {
        "DISEASE": [],
        "MEDICINE": [],
        "TREATMENT": [],
        "TEST": [],
        "ANATOMY": [],
        "SYMPTOMS": []
    }
    
    for result in results_list:
        if not result:
            continue
        for category in merged:
            if category in result:
                merged[category].extend(result[category])
    
    # Remove duplicates while preserving order
    for category in merged:
        seen = set()
        merged[category] = [x for x in merged[category] if not (x in seen or seen.add(x))]
    
    return merged

# Function to parse JSON from extracted text
def json_from_text(text):
    try:
        # Try parsing directly first
        return json.loads(text)
    except json.JSONDecodeError:
        # If that fails, try to find JSON in the text
        json_match = re.search(r'\{.*\}', text, re.DOTALL)
        if json_match:
            try:
                return json.loads(json_match.group())
            except:
                st.error("Found JSON-like content but couldn't parse it.")
                st.code(json_match.group(), language="json")
                return None
        else:
            st.error("No JSON found in response.")
            st.code(text)
            return None

# Function to translate text
def translate_text(data, target_language="en"):
    translator = GoogleTranslator(source='auto', target=target_language)
    
    translated_data = {}
    for category, items in data.items():
        translated_items = []
        with st.status(f"Translating {category}..."):
            for item in items:
                try:
                    # Check if text is too long (deep_translator has character limits)
                    if len(item) > 5000:
                        # Handle long text or just keep original
                        translated_items.append(item)
                        continue
                        
                    translated = translator.translate(item)
                    translated_items.append(translated)
                except Exception as e:
                    st.warning(f"Could not translate '{item}': {str(e)}")
                    translated_items.append(item)  # Keep original if translation fails
        
        translated_data[category] = translated_items
    
    return translated_data

# Function to create PDF report
def create_pdf_report(data):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    elements = []
    
    # Add title
    elements.append(Paragraph("Medical Data Extraction Report", styles['Title']))
    elements.append(Spacer(1, 12))
    
    # Add each category
    for category, items in data.items():
        if items:
            elements.append(Paragraph(f"{category} ({len(items)})", styles['Heading2']))
            elements.append(Spacer(1, 6))
            
            # Create table for items
            data_rows = [[item] for item in items]
            if data_rows:
                t = Table(data_rows, colWidths=[450])
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), colors.white),
                    ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ('GRID', (0, 0), (-1, -1), 1, colors.lightgrey)
                ]))
                elements.append(t)
            
            elements.append(Spacer(1, 12))
    
    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    return buffer

# Function to create DOCX report
def create_docx_report(data):
    doc = Document()
    doc.add_heading('Medical Data Extraction Report', 0)
    
    for category, items in data.items():
        if items:
            doc.add_heading(f"{category} ({len(items)})", level=2)
            
            # Add table
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Extracted Entities'
            
            for item in items:
                row_cells = table.add_row().cells
                row_cells[0].text = item
            
            doc.add_paragraph('')
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Chat function
def process_chat_message(user_message):
    try:
        response = llm.chat.completions.create(
            messages=[
                {"role": "system", "content": "You are a helpful medical assistant that answers questions about medical terminology, conditions, treatments, and medications. Keep your responses concise and informative."},
                {"role": "user", "content": user_message}
            ],
            model="mixtral-8x7b-32768"
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"I'm sorry, I encountered an error: {str(e)}"

# Dictionary of supported languages
LANGUAGES = {
    "en": "English",
    "zh-tw": "Chinese",
    "es": "Spanish",
    "hi": "Hindi",
    "ar": "Arabic"
}

# Initialize session state for chatbot
if "chat_messages" not in st.session_state:
    st.session_state.chat_messages = [
        {"role": "assistant", "content": "Hello! I'm your medical assistant. How can I help you today?"}
    ]

if "chat_visibility" not in st.session_state:
    st.session_state.chat_visibility = False

# Toggle chat visibility
def toggle_chat():
    st.session_state.chat_visibility = not st.session_state.chat_visibility

# Streamlit UI
st.set_page_config(layout="wide")

# st.markdown("""
#     <style>
#             .title {
#                 color: '#FF5733';
#                 text-align: center;
#             }
#     </style>
# """, unsafe_allow_html=True)
st.title("MEDIPARSER - A Medical Data Extractor")
# original_title = '<h2 style="font-family:Poppins,sans-serif; color: #38b2b5; font-size: 40px;">MEDIPARSER - A Medical Data Extractor</h2>'
# st.markdown(original_title, unsafe_allow_html=True)
# st.markdown('<div class="title"> Medical Data Extractor </div>', unsafe_allow_html=True)

# header = '<h3 style="font-family:Poppins,sans-serif; color: #E3D2C3    ; font-size: 25px; font-style: italic;">Upload medical records to extract structured information</h3>'
# st.markdown(header, unsafe_allow_html=True)

st.write("Upload patient medical records to extract structured information.")

# Sidebar for settings
with st.sidebar:
    st.header("Settings")
    
    # Translation options
    st.subheader("Translation")
    translate_enabled = st.checkbox("Enable translation", value=False)
    
    target_language = "en"
    if translate_enabled:
        language_options = {v: k for k, v in LANGUAGES.items()}
        language_names = list(language_options.keys())
        language_names.sort()
        selected_language = st.selectbox(
            "Target language", 
            options=language_names,
            index=language_names.index("english") if "english" in language_names else 0
        )
        target_language = language_options[selected_language]
    
    # Export options
    st.subheader("Export")
    export_format = st.selectbox("Export format", ["JSON", "CSV", "PDF", "DOCX"])

# File uploader for PDF, DOCX and text files
uploaded_file = st.file_uploader("Choose a file", type=["pdf", "txt", "docx"])

if uploaded_file:
    process_button = st.button("Extract Medical Data")
    
    if process_button:
        try:
            progress_bar = st.progress(0)
            
            # Extract text from file
            progress_bar.progress(10)
            st.info("Reading file...")
            
            if uploaded_file.type == "application/pdf":
                text = text_from_pdf(uploaded_file)
            elif uploaded_file.type == "text/plain":
                text = text_from_text(uploaded_file)
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                text = text_from_docx(uploaded_file)
            else:
                st.error("Unsupported file type")
                st.stop()
            
            # Show text preview
            progress_bar.progress(30)
            with st.expander("Document Text Preview"):
                st.write(text[:1000] + ("..." if len(text) > 1000 else ""))
            
            # Extract structured data
            progress_bar.progress(40)
            st.info("Analyzing medical data...")
            
            extracted_text = data_extraction(text)
            data = json_from_text(extracted_text)
            
            progress_bar.progress(80)
            
            if data:
                # Translate if enabled
                if translate_enabled:
                    st.info(f"Translating to {target_language}...")
                    translated_data = translate_text(data, target_language)
                    data_to_show = translated_data
                else:
                    data_to_show = data
                
                # Display results
                progress_bar.progress(100)
                
                st.success("Data extraction completed!")
                
                # Show results
                st.header("Extracted Medical Data")
                
                # Create tabs for different views
                tab1, tab2 = st.tabs(["Categorized View", "Raw JSON"])
                
                with tab1:
                    for category, items in data_to_show.items():
                        if items:
                            with st.expander(f"{category} ({len(items)})"):
                                for item in items:
                                    st.write(f"- {item}")
                
                with tab2:
                    st.json(data_to_show)
                
                # Export options
                if export_format == "JSON":
                    st.download_button(
                        "Download JSON",
                        data=json.dumps(data_to_show, indent=2),
                        file_name="medical_data.json",
                        mime="application/json"
                    )
                elif export_format == "CSV":
                    # Create CSV
                    csv_data = "Category,Entity\n"
                    for category, items in data_to_show.items():
                        for item in items:
                            csv_data += f"{category},{item.replace(',', ' ')}\n"
                    
                    st.download_button(
                        "Download CSV",
                        data=csv_data,
                        file_name="medical_data.csv",
                        mime="text/csv"
                    )
                elif export_format == "PDF":
                    pdf_buffer = create_pdf_report(data_to_show)
                    st.download_button(
                        "Download PDF",
                        data=pdf_buffer,
                        file_name="medical_data.pdf",
                        mime="application/pdf"
                    )
                elif export_format == "DOCX":
                    docx_buffer = create_docx_report(data_to_show)
                    st.download_button(
                        "Download DOCX",
                        data=docx_buffer,
                        file_name="medical_data.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            else:
                st.error("Failed to extract structured data. Please try again.")
        
        except Exception as e:
            st.error(f"An error occurred: {str(e)}")



# Chat toggle button at the bottom of the page
st.markdown("""
    <style>
    .chat-button {
        position: fixed;
        bottom: 20px;
        right: 20px;
        z-index: 9999;
        color: blue;
    }
    
    .chat-button > button {
        background-color: #4CAF50; /* Green background */
        color: white; /* White text */
        font-size: 16px; /* Font size */
        border-radius: 10px; /* Rounded corners */
        padding: 10px 20px; /* Padding */
        border: none; /* Remove border */
        cursor: pointer; /* Pointer cursor */
        box-shadow: 2px 2px 10px rgba(0, 0, 0, 0.2); /* Button shadow */
    }

    .chat-button > button:hover {
        background-color: #45a049; /* Darker green on hover */
    }
    </style>
""", unsafe_allow_html=True)

# Add button inside a div
st.markdown('<div class="chat-button">', unsafe_allow_html=True)
st.button("🤖 Chat with Medical Assistant", on_click=toggle_chat)
st.markdown('</div>', unsafe_allow_html=True)


# Chat container at the bottom
if st.session_state.chat_visibility:
    # Add custom styling for the chat container
    st.markdown("""
    <style>
    .chat-container {
        margin-top: 1px;
        border: 1px solid #e0e0e0;
        border-radius: 100px;
        padding: 15px;
        background-color: #f9f9f9;
    }
    .stChatMessage {
        padding: 10px;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Create a container for the chat
    chat_container = st.container()
    
    with chat_container:
        
        st.subheader("Medical Assistant Chat")       
        # Create columns for the chat interface
        chat_col1, chat_col2 = st.columns([5, 1])

        # Chat input
        if prompt := st.chat_input("Ask me about medical terms, conditions, or extracted data..."):
            # Add user message to chat history
            st.session_state.chat_messages.append({"role": "user", "content": prompt})
            
            # Display user message
            with st.chat_message("user"):
                st.write(prompt)
            
            # Display assistant response
            with st.chat_message("assistant"):
                message_placeholder = st.empty()
                message_placeholder.markdown("Thinking...")
                
                # Get response from LLM
                response = process_chat_message(prompt)
                
                # Update placeholder with response
                message_placeholder.markdown(response)
            
            # Add assistant response to chat history
            st.session_state.chat_messages.append({"role": "assistant", "content": response})
         
        # Display chat messages
        # for message in st.session_state.chat_messages:
        #     with st.chat_message(message["role"]):
        #         st.write(message["content"])
        
        
        st.markdown('</div>', unsafe_allow_html=True)
